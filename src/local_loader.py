"""
local_loader.py
───────────────
로컬 문서 기반 FAQ/RAG 소스 로더.

지원 형식:
- .md / .txt: 기본 지원
- .pdf: pypdf 또는 PyMuPDF(fitz)가 설치된 경우 지원
- .docx: python-docx가 설치된 경우 지원

목적:
- Notion 연결 없이 로컬 파일 업로드만으로 FAQ 지식베이스를 만들기
- 성형외과 FAQ 규정집처럼 `## FAQ-001. 질문` 구조의 문서를 Q/A 단위로 파싱
- 일반 문서는 섹션 단위 청킹으로 fallback
"""

from __future__ import annotations

import re
import shutil
from pathlib import Path
from typing import BinaryIO

BASE_DIR = Path(__file__).resolve().parents[1]
UPLOAD_DIR = BASE_DIR / "data" / "uploads"
SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx"}


def ensure_upload_dir() -> Path:
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    return UPLOAD_DIR


def save_uploaded_file(uploaded_file, upload_dir: Path | None = None) -> Path:
    """Streamlit uploaded_file 객체를 data/uploads에 저장하고 경로를 반환한다."""
    upload_dir = upload_dir or ensure_upload_dir()
    safe_name = sanitize_filename(uploaded_file.name)
    target = upload_dir / safe_name
    with target.open("wb") as f:
        f.write(uploaded_file.getbuffer())
    return target


def sanitize_filename(name: str) -> str:
    name = Path(name).name
    name = re.sub(r"[^0-9A-Za-z가-힣._ -]+", "_", name)
    return name.strip() or "uploaded_document"


def extract_text_from_file(path: str | Path) -> str:
    path = Path(path)
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"지원하지 않는 파일 형식입니다: {ext}")

    if ext in {".md", ".txt"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if ext == ".pdf":
        return extract_pdf_text(path)
    if ext == ".docx":
        return extract_docx_text(path)
    raise ValueError(f"지원하지 않는 파일 형식입니다: {ext}")


def extract_pdf_text(path: Path) -> str:
    """PDF 텍스트 추출. pypdf 우선, 실패 시 PyMuPDF 시도."""
    try:
        from pypdf import PdfReader  # type: ignore
        reader = PdfReader(str(path))
        pages = []
        for i, page in enumerate(reader.pages, 1):
            text = page.extract_text() or ""
            pages.append(f"\n\n[PDF Page {i}]\n{text}")
        return "".join(pages)
    except ImportError:
        pass

    try:
        import fitz  # type: ignore
        doc = fitz.open(str(path))
        pages = []
        for i, page in enumerate(doc, 1):
            pages.append(f"\n\n[PDF Page {i}]\n{page.get_text()}")
        return "".join(pages)
    except ImportError as e:
        raise RuntimeError("PDF 지원을 위해 `pypdf` 또는 `pymupdf` 설치가 필요합니다.") from e


def extract_docx_text(path: Path) -> str:
    try:
        import docx  # type: ignore
    except ImportError as e:
        raise RuntimeError("Word(.docx) 지원을 위해 `python-docx` 설치가 필요합니다.") from e

    doc = docx.Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def load_faq_from_files(paths: list[str | Path]) -> list[dict]:
    """여러 로컬 문서를 읽어 FAQ 리스트로 변환한다."""
    all_items: list[dict] = []
    for path in paths:
        p = Path(path)
        text = extract_text_from_file(p)
        items = parse_faq_text(text, source_name=p.name)
        all_items.extend(items)
    return all_items


def parse_faq_text(text: str, source_name: str = "local_document") -> list[dict]:
    """FAQ 형식 문서는 FAQ 단위, 일반 문서는 섹션 단위로 파싱한다."""
    faq_items = parse_structured_faq(text, source_name=source_name)
    if faq_items:
        return faq_items
    return parse_section_chunks(text, source_name=source_name)


def parse_structured_faq(text: str, source_name: str = "local_document") -> list[dict]:
    """`## FAQ-001. 질문` 형태의 문서를 Q/A 단위로 파싱한다."""
    pattern = re.compile(
        r"(?ms)^#{1,3}\s*(FAQ[- ]?\d+)\.\s*(.+?)\n(.*?)(?=^#{1,3}\s*FAQ[- ]?\d+\.|\Z)"
    )
    items: list[dict] = []
    for match in pattern.finditer(text):
        faq_id = match.group(1).strip()
        title = clean_markdown(match.group(2).strip())
        body = match.group(3).strip()
        answer = extract_answer_block(body) or body
        examples = extract_question_examples(body)
        tags = extract_tag_block(body)
        question = title
        if examples:
            question = f"{title}\n질문 예시: " + " / ".join(examples[:5])
        items.append({
            "question": question,
            "answer": clean_answer(answer),
            "source": source_name,
            "faq_id": faq_id,
            "tags": ", ".join(tags),
            "chunk_type": "faq",
        })
    return items


def extract_answer_block(body: str) -> str:
    markers = ["**답변**", "답변", "### 답변"]
    start = -1
    for marker in markers:
        start = body.find(marker)
        if start >= 0:
            start += len(marker)
            break
    if start < 0:
        return ""

    tail = body[start:]
    # 다음 굵은 섹션/구분선 전까지
    stop_match = re.search(r"(?m)^\*\*[^\n]+\*\*\s*$|^---\s*$", tail)
    if stop_match and stop_match.start() > 0:
        tail = tail[:stop_match.start()]
    return tail.strip(" \n:")


def extract_question_examples(body: str) -> list[str]:
    match = re.search(r"(?ms)\*\*질문 예시\*\*\s*(.*?)(?=^\*\*답변\*\*|^---|\Z)", body)
    if not match:
        return []
    block = match.group(1)
    examples = []
    for line in block.splitlines():
        line = line.strip().lstrip("-• ").strip()
        if line:
            examples.append(clean_markdown(line))
    return examples


def extract_tag_block(body: str) -> list[str]:
    match = re.search(r"(?ms)\*\*태그\*\*\s*(.*?)(?=^\*\*|^---|\Z)", body)
    if not match:
        return []
    raw = clean_markdown(match.group(1))
    return [t.strip() for t in re.split(r"[,，\n]", raw) if t.strip()]


def parse_section_chunks(text: str, source_name: str = "local_document") -> list[dict]:
    """FAQ 패턴이 없으면 Markdown 제목 기준으로 섹션 청킹한다."""
    sections = re.split(r"(?m)^(#{1,3}\s+.+)$", text)
    items: list[dict] = []
    if len(sections) <= 1:
        chunks = chunk_text(text, max_chars=1200)
        for i, chunk in enumerate(chunks, 1):
            items.append({
                "question": f"{source_name} 문서 내용 {i}",
                "answer": clean_answer(chunk),
                "source": source_name,
                "faq_id": f"CHUNK-{i:03d}",
                "tags": "",
                "chunk_type": "section",
            })
        return items

    prefix = sections[0].strip()
    for i in range(1, len(sections), 2):
        heading = clean_markdown(sections[i].strip("# ").strip())
        content = sections[i + 1].strip() if i + 1 < len(sections) else ""
        if len(content) < 30:
            continue
        for j, chunk in enumerate(chunk_text(content, max_chars=1400), 1):
            suffix = f" #{j}" if j > 1 else ""
            items.append({
                "question": f"{heading}{suffix}",
                "answer": clean_answer(chunk),
                "source": source_name,
                "faq_id": f"SEC-{len(items)+1:03d}",
                "tags": "",
                "chunk_type": "section",
            })
    return items


def chunk_text(text: str, max_chars: int = 1200) -> list[str]:
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    chunks: list[str] = []
    current = ""
    for para in paragraphs:
        if len(current) + len(para) + 2 <= max_chars:
            current = f"{current}\n\n{para}".strip()
        else:
            if current:
                chunks.append(current)
            current = para
    if current:
        chunks.append(current)
    return chunks


def clean_markdown(text: str) -> str:
    text = re.sub(r"[`*_>#]", "", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def clean_answer(text: str) -> str:
    text = text.replace("\r\n", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


if __name__ == "__main__":
    sample = BASE_DIR.parent / "plastic_surgery_faq_knowledge_base.md"
    if sample.exists():
        faqs = load_faq_from_files([sample])
        print(f"loaded_faq_count={len(faqs)}")
        for item in faqs[:3]:
            print("---")
            print(item["faq_id"], item["question"][:120])
            print(item["answer"][:160])
    else:
        print("샘플 파일을 찾지 못했습니다.")
